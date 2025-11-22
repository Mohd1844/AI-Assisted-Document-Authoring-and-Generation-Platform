from docx import Document
from docx.shared import Pt, Inches
from typing import List
import io

def create_docx(title: str, sections: List[dict]) -> bytes:
    """
    Create a Word document
    sections: [{'title': 'Section 1', 'content': 'Content here'}, ...]
    """
    doc = Document()
    
    # Add title
    title_paragraph = doc.add_heading(title, level=0)
    title_paragraph.alignment = 1  # Center alignment
    
    # Add sections
    for section in sections:
        doc.add_heading(section['title'], level=1)
        doc.add_paragraph(section['content'])
        doc.add_paragraph()  # Add spacing
    
    # Save to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.read()
